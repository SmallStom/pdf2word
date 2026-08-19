# -*- coding: utf-8 -*-
"""Word文档生成模块 - 按照固定格式规范生成Word文档

格式规范：
- 页面：A4，上下2.54cm，左右3.18cm，页眉1.5cm，页脚1.75cm
- 正文：宋体+Times New Roman，小四号(12pt)，1.5倍行距，首行缩进2字符
- 标题：宋体加粗，一级三号(16pt)/二级小三号(15pt)/三级四号(14pt)/四级+小四号(12pt)
- 表格：居中，行高0.75cm，表内五号(10.5pt)单倍行距，表名小四加粗居中在上方
- 图片：居中嵌入型，图名居中在下方，分章编号
"""

import re
import io
from typing import List, Optional

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import (
    PAGE_SIZE, PAGE_MARGINS, HEADER_FOOTER_DISTANCE,
    HEADING_SIZES, HEADING_FONT, BODY_FONT,
    TABLE_CONFIG, IMAGE_CONFIG, USABLE_WIDTH_CM,
)
from pdf_extractor import ContentElement
from style_mapper import FigureTableCounter


# ============================================================
# 页眉页脚检测
# ============================================================
_HEADER_FOOTER_RE = re.compile(
    r'^(第\s*\d+\s*页|'
    r'-\s*\d+\s*-|'                  # - 5 -
    r'第\s*[零一二三四五六七八九十百千0-9]+\s*页|'
    r'Page\s*\d+|'                   # Page 3
    r'\d+\s*/\s*\d+\s*$|'            # 3/10
    r'[\-—]\s*\d+\s*[\-—]\s*$)'      # -3- / —3—
)


def _is_likely_header_footer(text: str) -> bool:
    """判断文本是否像页眉页脚

    规则：
    1. 长度 <= 20 字符
    2. 匹配"第N页"/"-3-"/"Page 3"/"3/10"等页码模式
    3. 全是数字或"-"的组合
    """
    text = text.strip()
    if len(text) > 20:
        return False
    if _HEADER_FOOTER_RE.match(text):
        return True
    # 纯数字/罗马数字短文本
    if len(text) <= 6 and re.fullmatch(r'[\dIVXLCDM\-—/、\s]+', text):
        return True
    return False


def _clean_text(text: str) -> str:
    """清理文本中的常见 OCR 错误和半全角混用

    - 统一换行格式：\r\n / \r → \n
    - 修复 https：// → https://
    - 合并多余半角空格
    - 修复 "第二章 投标人须知" 这种"中文标题+半角空格+中文"
    """
    if not text:
        return text
    text = text.strip()

    # 统一换行格式
    text = re.sub(r'\r\n?', '\n', text)

    # 修复全角冒号/斜杠（在 URL 场景）
    text = re.sub(r'(https?|ftp)\s*[：:]\s*/\s*/\s*', r'\1://', text)
    text = re.sub(r'([：:])\s*/\s*/\s*', r'://', text)

    # "1.1" 后面多余空格
    text = re.sub(r'^(\d+(?:\.\d+)+)\s{2,}', r'\1 ', text)

    # 修复连续标点
    text = re.sub(r'，\s*，', '，', text)
    text = re.sub(r'。\s*。', '。', text)

    # 统一内部换行：连续 \n 压缩，行首行尾去空白
    text = re.sub(r'\n\s*\n+', '\n', text)
    text = '\n'.join(line.strip() for line in text.split('\n'))

    return text.strip()


# ============================================================
# 底层XML操作工具
# ============================================================
def _set_run_font(run, cn_font='宋体', en_font='Times New Roman',
                  size_pt=12, bold=False, italic=False):
    """设置Run的字体，正确处理中文（w:eastAsia / w:ascii / w:hAnsi）

    注意：必须同时设置 ascii/hAnsi/eastAsia/cs 四种字体属性，
    不能只设置 run.font.name（python-docx 会清空 eastAsia）
    """
    # 1. 通过 XML 直接设置所有字体属性（关键！）
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    # 必须同时设置 4 个属性，Word 才能正确显示中文
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:cs'), en_font)

    # 2. 字号、加粗、斜体
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def _set_paragraph_spacing(paragraph, line_spacing=1.5,
                            space_before=0, space_after=0,
                            snap_to_grid=False):
    """设置段落间距：行距、段前段后、网格对齐"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)

    # 取消"如果定义了文档网格，则对齐到网格"
    if not snap_to_grid:
        pPr = paragraph._element.get_or_add_pPr()
        snapToGrid = pPr.find(qn('w:snapToGrid'))
        if snapToGrid is None:
            snapToGrid = OxmlElement('w:snapToGrid')
            pPr.append(snapToGrid)
        snapToGrid.set(qn('w:val'), '0')


def _set_row_height(row, height_cm, rule='atLeast'):
    """设置表格行高"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    # 1cm = 567 twips
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), rule)
    trPr.append(trHeight)


def _set_repeat_header(row):
    """设置表格第一行为重复标题行（跨页时重复）"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def _set_cant_split(row):
    """设置表格行不可跨页拆分"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    cantSplit.set(qn('w:val'), 'true')
    trPr.append(cantSplit)


def _set_table_borders(table):
    """设置表格边框（黑色实线）"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        borders.append(border)
    tblPr.append(borders)


# ============================================================
# 页面设置
# ============================================================
def setup_page(doc: Document):
    """设置A4页面、页边距、页眉页脚距离"""
    section = doc.sections[0]
    section.page_width = Cm(PAGE_SIZE['width_cm'])
    section.page_height = Cm(PAGE_SIZE['height_cm'])
    section.top_margin = Cm(PAGE_MARGINS['top_cm'])
    section.bottom_margin = Cm(PAGE_MARGINS['bottom_cm'])
    section.left_margin = Cm(PAGE_MARGINS['left_cm'])
    section.right_margin = Cm(PAGE_MARGINS['right_cm'])
    section.header_distance = Cm(HEADER_FOOTER_DISTANCE['header_cm'])
    section.footer_distance = Cm(HEADER_FOOTER_DISTANCE['footer_cm'])


def setup_page_number_footer(doc: Document, start_at: int = 1):
    """在 footer 添加 "第 X 页 共 Y 页" 格式的页码

    启动编号从 start_at 开始（封面页可不计入）
    """
    section = doc.sections[0]
    footer = section.footer
    # 清空已有段落
    for p in list(footer.paragraphs):
        if p.text:
            for r in p.runs:
                r.text = ''

    # 创建页码段落
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ''
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "第 "
    run1 = p.add_run('第 ')
    _set_run_font(run1, cn_font='宋体', en_font='Times New Roman', size_pt=10.5)

    # PAGE 字段（当前页）
    _add_field(p, 'PAGE')

    # " 页 共 "
    run2 = p.add_run(' 页 共 ')
    _set_run_font(run2, cn_font='宋体', en_font='Times New Roman', size_pt=10.5)

    # NUMPAGES 字段（总页数）
    _add_field(p, 'NUMPAGES')

    # " 页"
    run3 = p.add_run(' 页')
    _set_run_font(run3, cn_font='宋体', en_font='Times New Roman', size_pt=10.5)

    # 启动编号
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start_at))


def _add_field(paragraph, field_code: str):
    """在段落里插入 Word 字段（如 PAGE / NUMPAGES）"""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' {field_code} '
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar_begin)
    run._element.append(instrText)
    run._element.append(fldChar_separate)
    run._element.append(fldChar_end)

    # 字段默认显示 "1"
    run.text = '1'
    _set_run_font(run, cn_font='宋体', en_font='Times New Roman', size_pt=10.5)


# ============================================================
# 正文段落
# ============================================================
def _add_text_runs(paragraph, runs: List[dict], default_size_pt: float,
                   bold: bool = False):
    """向段落写入多个run（保留PDF行内异体样式：字号/加粗/颜色）

    runs 每项: {'text', 'size', 'bold', 'italic', 'color'}
    文本内的换行符转为 Word 软换行（<w:br/>）。
    """
    for rd in runs:
        text = rd.get('text') or ''
        if not text:
            continue
        size = rd.get('size') or default_size_pt
        run_bold = bool(rd.get('bold', bold))
        segments = text.split('\n')
        run = None
        for i, seg in enumerate(segments):
            if i > 0:
                run.add_break()
            if seg:
                run = paragraph.add_run(seg)
                _set_run_font(run,
                              cn_font=BODY_FONT['cn'],
                              en_font=BODY_FONT['en'],
                              size_pt=size,
                              bold=run_bold,
                              italic=bool(rd.get('italic', False)))
                color = rd.get('color')
                if color is not None and int(color) != 0:
                    c = int(color)
                    run.font.color.rgb = RGBColor(
                        (c >> 16) & 0xFF, (c >> 8) & 0xFF, c & 0xFF)
        # 段末换行：补一个换行run（极少出现）
        if text.endswith('\n') and run is not None:
            run.add_break()


def add_body_paragraph(doc: Document, text: str, size_pt: float = None,
                       alignment: str = 'left', bold: bool = False,
                       runs: Optional[List[dict]] = None,
                       color: Optional[int] = None,
                       left_indent_pt: float = 0.0,
                       first_line_indent_pt: Optional[float] = None):
    """添加正文段落

    宋体+Times New Roman，小四号(12pt)或指定字号，1.5倍行距，首行缩进2字符
    - runs: PDF行内异体样式run列表（保留混排字号/加粗/颜色）
    - color: 段落主色（runs为空时生效）
    - left_indent_pt: 左缩进（目录层级/块缩进还原）
    - first_line_indent_pt: 首行缩进；None=配置的2字符，
      数值=按PDF几何还原（悬挂/深缩进，已按字号映射缩放）
    """
    p = doc.add_paragraph()
    actual_size = size_pt if size_pt and size_pt > 0 else BODY_FONT['size_pt']

    if runs:
        _add_text_runs(p, runs, actual_size, bold=bold)
    else:
        _add_text_runs(p, [{'text': text, 'size': actual_size,
                            'bold': bold, 'italic': False,
                            'color': color}],
                       actual_size, bold=bold)

    _set_paragraph_spacing(p,
                           line_spacing=BODY_FONT['line_spacing'],
                           space_before=BODY_FONT['space_before_pt'],
                           space_after=BODY_FONT['space_after_pt'],
                           snap_to_grid=BODY_FONT['snap_to_grid'])
    # 首行缩进（仅左对齐正文需要，居中/右对齐不缩进）
    # None/未指定 = 配置的2字符；数值 = PDF几何还原
    if alignment == 'left' and not _is_toc_entry(text):
        if first_line_indent_pt is not None:
            p.paragraph_format.first_line_indent = Pt(max(0.0, first_line_indent_pt))
        else:
            indent_pt = actual_size * BODY_FONT['first_line_indent_chars']
            p.paragraph_format.first_line_indent = Pt(indent_pt)
    if left_indent_pt > 0:
        p.paragraph_format.left_indent = Pt(left_indent_pt)
    if alignment == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _is_toc_entry(text: str) -> bool:
    """判断是否为目录条目（含前导点号连接页码）"""
    if not text:
        return False
    # 匹配 "... 12" 或 "…… 12" 或 "..82"
    return bool(re.search(r'[\.\.]{2,}\s*\d+\s*$', text) or
                re.search(r'[…⋯]{2,}\s*\d+\s*$', text))


# ============================================================
# 目录条目段落（标题 + 制表位引导符 + 右对齐页码）
# ============================================================
def add_toc_paragraph(doc: Document, title: str, page_no: str,
                      left_indent_pt: float = 0.0,
                      size_pt: float = None,
                      usable_width_cm: float = None):
    """重建目录条目：标题 + 右对齐制表位（点状引导符）+ 页码

    结构与Word原生目录一致：
    - 制表位定位在可用宽度右端，前导符为点线
    - 页码自动右对齐到行尾
    - 左缩进保留原文层级
    """
    p = doc.add_paragraph()
    actual_size = size_pt if size_pt and size_pt > 0 else BODY_FONT['size_pt']
    width_cm = usable_width_cm or USABLE_WIDTH_CM
    # 制表位位置 = 可用宽度 - 左缩进（相对缩进后的文本区）
    tab_pos_cm = max(1.0, width_cm - left_indent_pt * 0.0352778)
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(tab_pos_cm), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    # 标题run
    run_title = p.add_run(title)
    _set_run_font(run_title, cn_font=BODY_FONT['cn'],
                  en_font=BODY_FONT['en'], size_pt=actual_size)
    # 制表符run
    run_tab = p.add_run()
    run_tab.add_tab()
    _set_run_font(run_tab, cn_font=BODY_FONT['cn'],
                  en_font=BODY_FONT['en'], size_pt=actual_size)
    # 页码run
    run_page = p.add_run(page_no)
    _set_run_font(run_page, cn_font=BODY_FONT['cn'],
                  en_font=BODY_FONT['en'], size_pt=actual_size)

    _set_paragraph_spacing(p,
                           line_spacing=BODY_FONT['line_spacing'],
                           space_before=0,
                           space_after=0,
                           snap_to_grid=False)
    if left_indent_pt > 0:
        p.paragraph_format.left_indent = Pt(left_indent_pt)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ============================================================
# 标题段落
# ============================================================
def add_heading(doc: Document, text: str, level: int, alignment: str = 'left',
                page_break_before: bool = False):
    """添加标题段落

    使用 Word 原生 Heading 1/2/3 样式（支持大纲/目录/导航）
    宋体加粗黑色，按层级设置字号（16/15/14/12pt）
    alignment: 'left' / 'center' / 'right'（按 PDF 实际位置推断）
    page_break_before: 是否在标题前另起一页（一级标题用）
    """
    # 使用 Word 原生 Heading 样式，让 Word 知道这是标题
    # 这样 Word 的大纲视图、目录生成、导航窗格都能识别
    heading_style_name = f'Heading {level}' if 1 <= level <= 9 else 'Heading 4'
    try:
        p = doc.add_paragraph(style=heading_style_name)
    except KeyError:
        # 退化：手动建样式
        p = doc.add_paragraph()

    # 一级标题另起一页
    if page_break_before:
        p.paragraph_format.page_break_before = True

    # 清空原有的样式字体（python-docx 默认 Heading 样式可能是 Calibri 蓝色）
    p.style.font.name = HEADING_FONT['en']
    p.style.font.size = Pt(HEADING_SIZES.get(level, HEADING_SIZES[4]))
    p.style.font.bold = HEADING_FONT['bold']
    # 关键：强制标题为黑色（Word内置Heading样式默认是蓝色主题色）
    p.style.font.color.rgb = RGBColor(0, 0, 0)
    # 同时设置 style 的中文/西文字体
    rPr_style = p.style.element.get_or_add_rPr()
    rFonts_style = rPr_style.find(qn('w:rFonts'))
    if rFonts_style is None:
        rFonts_style = OxmlElement('w:rFonts')
        rPr_style.insert(0, rFonts_style)
    rFonts_style.set(qn('w:ascii'), HEADING_FONT['en'])
    rFonts_style.set(qn('w:hAnsi'), HEADING_FONT['en'])
    rFonts_style.set(qn('w:eastAsia'), HEADING_FONT['cn'])
    rFonts_style.set(qn('w:cs'), HEADING_FONT['en'])

    # 添加 run 并设置 run 自己的字体（覆盖 style 的）
    text = _clean_text(text)
    size = HEADING_SIZES.get(level, HEADING_SIZES[4])
    for seg in text.split('\n'):
        if not seg:
            continue
        run = p.add_run(seg)
        _set_run_font(run,
                      cn_font=HEADING_FONT['cn'],
                      en_font=HEADING_FONT['en'],
                      size_pt=size,
                      bold=HEADING_FONT['bold'])
        run.font.color.rgb = RGBColor(0, 0, 0)
    _set_paragraph_spacing(p,
                           line_spacing=BODY_FONT['line_spacing'],
                           space_before=0,
                           space_after=0,
                           snap_to_grid=False)
    # 标题对齐：按 PDF 实际位置
    if alignment == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ============================================================
# 表格生成
# ============================================================
def _parse_html_table(html: str) -> List[List[str]]:
    """解析HTML表格，返回行列数据

    支持 colspan/rowspan/嵌套表格，规范化每个 row 的列数
    """
    rows = []
    # 1. 先把整段 HTML 里 <tbody>/<thead>/<tfoot> 包起来
    tr_pattern = re.compile(r'<tr[^>]*>(.*?)</tr>', re.DOTALL | re.IGNORECASE)
    td_pattern = re.compile(
        r'<(t[dh])\b([^>]*)>(.*?)</\1>',
        re.DOTALL | re.IGNORECASE
    )

    for tr_match in tr_pattern.finditer(html):
        row = []
        for td_match in td_pattern.finditer(tr_match.group(1)):
            tag = td_match.group(1).lower()
            attrs = td_match.group(2)
            cell_text = td_match.group(3)

            # 提取 colspan/rowspan
            colspan = 1
            rowspan = 1
            m = re.search(r'colspan\s*=\s*["\']?(\d+)["\']?', attrs, re.IGNORECASE)
            if m:
                colspan = max(int(m.group(1)), 1)
            m = re.search(r'rowspan\s*=\s*["\']?(\d+)["\']?', attrs, re.IGNORECASE)
            if m:
                rowspan = max(int(m.group(1)), 1)

            # 去除嵌套表格：递归提取子表格的文本内容
            inner_table = re.search(r'<table\b.*?</table>', cell_text, re.DOTALL | re.IGNORECASE)
            if inner_table:
                # 提取子表所有 td 的文字拼到当前单元格
                inner_texts = re.findall(
                    r'<t[dh][^>]*>(.*?)</t[dh]>',
                    inner_table.group(0),
                    re.DOTALL | re.IGNORECASE
                )
                cell_text = '\n'.join(re.sub(r'<[^>]+>', '', t) for t in inner_texts)
            else:
                # 去除嵌套HTML标签
                cell_text = re.sub(r'<[^>]+>', '', cell_text)

            # 解码HTML实体
            cell_text = cell_text.replace('&lt;', '<').replace('&gt;', '>')
            cell_text = cell_text.replace('&amp;', '&').replace('&nbsp;', ' ')
            cell_text = cell_text.replace('&quot;', '"').replace('&#39;', "'")
            cell_text = cell_text.strip()
            # 把多空白压缩成单个换行（保留段落感）
            cell_text = re.sub(r'\s*\n\s*', '\n', cell_text)

            # 按 colspan 展开
            for _ in range(colspan):
                row.append(cell_text if colspan == 1 else '')
            # 如果 colspan > 1，第一个 cell 保留文字，其余为空
            if colspan > 1:
                row[-colspan] = cell_text
        if row:
            rows.append(row)
    return rows


def add_table(doc: Document, html: str, counter: FigureTableCounter):
    """添加表格

    - 表名在上方，居中，宋体小四号加粗
    - 表格居中，无文字环绕
    - 行高0.75cm
    - 表内文字宋体五号(10.5pt)，单倍行距
    - 重复标题行
    - 行不跨页拆分
    - 表名与表格同页
    """
    rows = _parse_html_table(html)
    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    if num_cols == 0:
        return

    # 规范化：所有行补齐到 num_cols
    for row in rows:
        while len(row) < num_cols:
            row.append('')

    # 1. 添加表名（在表格上方）
    table_name = counter.next_table_number()
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(table_name)
    _set_run_font(name_run,
                  cn_font=TABLE_CONFIG['name_font_cn'],
                  en_font=TABLE_CONFIG['name_font_en'],
                  size_pt=TABLE_CONFIG['name_font_size_pt'],
                  bold=TABLE_CONFIG['name_bold'])
    _set_paragraph_spacing(name_paragraph,
                           line_spacing=BODY_FONT['line_spacing'],
                           space_before=0,
                           space_after=0,
                           snap_to_grid=False)
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 表名与表格保持在同一页
    name_paragraph.paragraph_format.keep_with_next = True

    # 2. 创建表格
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    # 3. 填充表格内容并设置样式
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        # 行高0.75cm
        _set_row_height(row, TABLE_CONFIG['row_height_cm'],
                        TABLE_CONFIG['row_height_rule'])
        # 行不跨页拆分
        if TABLE_CONFIG['cant_split_row']:
            _set_cant_split(row)
        # 第一行设为重复标题行
        if i == 0 and TABLE_CONFIG['repeat_header']:
            _set_repeat_header(row)

        for j in range(num_cols):
            cell_text = row_data[j] if j < len(row_data) else ''
            cell = row.cells[j]
            # 清除默认内容
            cell.text = ''
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(cell_text)
            _set_run_font(run,
                          cn_font=TABLE_CONFIG['cell_font_cn'],
                          en_font=TABLE_CONFIG['cell_font_en'],
                          size_pt=TABLE_CONFIG['cell_font_size_pt'],
                          bold=False)
            _set_paragraph_spacing(paragraph,
                                   line_spacing=TABLE_CONFIG['cell_line_spacing'],
                                   space_before=TABLE_CONFIG['cell_space_before_pt'],
                                   space_after=TABLE_CONFIG['cell_space_after_pt'],
                                   snap_to_grid=False)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


# ============================================================
# 图片插入
# ============================================================
def add_image(doc: Document, image_data: bytes, counter: FigureTableCounter):
    """添加图片

    - 嵌入型（inline），居中
    - 图名在下方，居中，分章编号
    - 图与名在同一页
    """
    if not image_data:
        return

    # 1. 插入图片（默认嵌入型）
    image_stream = io.BytesIO(image_data)
    try:
        doc.add_picture(image_stream, width=Cm(IMAGE_CONFIG['max_width_cm']))
    except Exception:
        return

    # 2. 设置图片段落居中，并与图名保持同页
    image_paragraph = doc.paragraphs[-1]
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True

    # 3. 添加图名（在图下方，居中，分章编号）
    figure_name = counter.next_figure_number()
    name_paragraph = doc.add_paragraph()
    name_run = name_paragraph.add_run(figure_name)
    _set_run_font(name_run,
                  cn_font=IMAGE_CONFIG['name_font_cn'],
                  en_font=IMAGE_CONFIG['name_font_en'],
                  size_pt=IMAGE_CONFIG['name_font_size_pt'],
                  bold=IMAGE_CONFIG['name_bold'])
    _set_paragraph_spacing(name_paragraph,
                           line_spacing=BODY_FONT['line_spacing'],
                           space_before=0,
                           space_after=0,
                           snap_to_grid=False)
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ============================================================
# 分节方向管理（横版/竖版页面）
# ============================================================
def _apply_section_page_setup(section, landscape: bool):
    """设置分节页面方向与尺寸（A4）"""
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(PAGE_SIZE['height_cm'])
        section.page_height = Cm(PAGE_SIZE['width_cm'])
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(PAGE_SIZE['width_cm'])
        section.page_height = Cm(PAGE_SIZE['height_cm'])
    section.top_margin = Cm(PAGE_MARGINS['top_cm'])
    section.bottom_margin = Cm(PAGE_MARGINS['bottom_cm'])
    section.left_margin = Cm(PAGE_MARGINS['left_cm'])
    section.right_margin = Cm(PAGE_MARGINS['right_cm'])
    section.header_distance = Cm(HEADER_FOOTER_DISTANCE['header_cm'])
    section.footer_distance = Cm(HEADER_FOOTER_DISTANCE['footer_cm'])


def _usable_width_cm(landscape: bool) -> float:
    if landscape:
        return PAGE_SIZE['height_cm'] - PAGE_MARGINS['left_cm'] - PAGE_MARGINS['right_cm']
    return USABLE_WIDTH_CM


# ============================================================
# 主生成函数
# ============================================================
def generate_word(elements: List[ContentElement], output_path: str):
    """从结构化数据生成Word文档

    遍历所有ContentElement，按类型生成对应内容：
    - heading -> 标题段落
    - toc -> 目录条目（制表位 + 点状引导符 + 右对齐页码）
    - text -> 正文段落
    - table -> 表格（含表名）
    - image -> 图片（含图名）
    横版页面的元素自动切换到横向分节。
    """
    doc = Document()

    # 1. 设置页面格式
    setup_page(doc)

    # 1.5 设置页码（第 X 页 共 Y 页）
    setup_page_number_footer(doc, start_at=1)

    # 2. 设置默认Normal样式（同时设置 ascii/hAnsi/eastAsia/cs 四种字体）
    style = doc.styles['Normal']
    style.font.name = BODY_FONT['en']
    style.font.size = Pt(BODY_FONT['size_pt'])
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), BODY_FONT['en'])
    rFonts.set(qn('w:hAnsi'), BODY_FONT['en'])
    rFonts.set(qn('w:eastAsia'), BODY_FONT['cn'])
    rFonts.set(qn('w:cs'), BODY_FONT['en'])

    # 3. 图表编号计数器
    counter = FigureTableCounter()

    # 4. 遍历内容元素
    import logging
    _log = logging.getLogger(__name__)
    added_counts = {'heading': 0, 'text': 0, 'toc': 0, 'table': 0, 'image': 0,
                    'skipped_empty': 0, 'skipped_header': 0}
    content_added = False  # 记录是否已有内容，决定一级标题是否另起一页
    current_landscape = False  # 当前分节方向（初始为竖版首页）
    for elem in elements:
        # 横版/竖版切换：另起新分节
        if elem.is_landscape != current_landscape:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            _apply_section_page_setup(section, elem.is_landscape)
            current_landscape = elem.is_landscape

        if elem.type == 'toc':
            # 目录条目：标题 + 制表位引导符 + 右对齐页码
            if elem.toc_title:
                add_toc_paragraph(
                    doc, elem.toc_title, elem.toc_page_no,
                    left_indent_pt=elem.left_indent_pt,
                    usable_width_cm=_usable_width_cm(current_landscape),
                )
                content_added = True
                added_counts['toc'] += 1

        elif elem.type == 'heading':
            # 更新章节号
            counter.update_chapter(elem.text, elem.heading_level or 1)
            # 添加标题（传 alignment 保持 PDF 原版位置）
            if elem.text and elem.text.strip():
                # 一级标题另起一页：已有正文才分页，避免文档开头空白页
                page_break = (elem.heading_level == 1) and content_added
                add_heading(doc, elem.text, elem.heading_level or 1,
                            alignment=elem.alignment or 'left',
                            page_break_before=page_break)
                content_added = True
                added_counts['heading'] += 1

        elif elem.type == 'table':
            # 添加表格
            if elem.html:
                add_table(doc, elem.html, counter)
                content_added = True
                added_counts['table'] += 1
            else:
                _log.warning(f"[DIAG-gen] 表格被跳过：无 html, text={elem.text[:50]!r}")

        elif elem.type == 'image':
            # 添加图片
            if elem.image_data:
                add_image(doc, elem.image_data, counter)
                content_added = True
                added_counts['image'] += 1
            else:
                _log.warning(f"[DIAG-gen] 图片被跳过：无 image_data, text={elem.text[:50]!r}")

        else:
            # 正文
            text = _clean_text(elem.text.strip())
            if not text:
                added_counts['skipped_empty'] += 1
                continue

            # 过滤明显是页眉页脚的内容：单独成行且极短、纯数字
            if _is_likely_header_footer(text):
                added_counts['skipped_header'] += 1
                continue

            add_body_paragraph(
                doc, text,
                size_pt=elem.mapped_size,
                alignment=elem.alignment or 'left',
                bold=bool(elem.is_bold),
                runs=elem.runs,
                color=elem.color,
                left_indent_pt=elem.left_indent_pt,
                first_line_indent_pt=elem.first_line_indent_pt,
            )
            content_added = True
            added_counts['text'] += 1

    _log.info(f"[DIAG-gen] 生成统计: {added_counts}")

    # 5. 保存文档
    doc.save(output_path)
