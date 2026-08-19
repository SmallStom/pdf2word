# -*- coding: utf-8 -*-
"""端到端测试：test.pdf -> docx，并dump文档结构验证"""
import sys, io, os
sys.path.insert(0, '.')
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from pipeline import PDFToWordPipeline

pdf = 'dataset/test.pdf'
out = 'output/test_e2e.docx'
os.makedirs('output', exist_ok=True)

pipe = PDFToWordPipeline()
result = pipe.process(pdf, out)
print('OK:', result, 'size=', os.path.getsize(result))

# 验证docx结构
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
doc = Document(out)

with open('scripts/out_docx.txt', 'w', encoding='utf-8') as f:
    f.write(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}\n\n")
    f.write("=== 前 45 个段落 ===\n")
    for i, p in enumerate(doc.paragraphs[:45]):
        style = p.style.name
        align = str(p.alignment).split('.')[-1].split(' ')[0] if p.alignment is not None else 'None'
        runs = len(p.runs)
        text = p.text[:70]
        tabs = len(p.paragraph_format.tab_stops)
        f.write(f"[{i:3d}] style={style:12s} align={align:8s} runs={runs} tabs={tabs} {text!r}\n")
    f.write("\n=== TOC段落（含制表位）===\n")
    cnt = 0
    for i, p in enumerate(doc.paragraphs):
        if p.paragraph_format.tab_stops and cnt < 8:
            f.write(f"[{i:3d}] tabs={len(p.paragraph_format.tab_stops)} {p.text[:60]!r}\n")
            cnt += 1
    f.write("\n=== 标题段落分布 ===\n")
    from collections import Counter
    c = Counter(p.style.name for p in doc.paragraphs)
    f.write(str(dict(c)) + '\n')
    f.write(f"\n=== 分节 ===\n")
    for i, s in enumerate(doc.sections):
        f.write(f"section {i}: orient={'LANDSCAPE' if s.orientation == 1 else 'portrait'} "
                f"w={s.page_width.cm:.1f}cm h={s.page_height.cm:.1f}cm\n")
print('docx dumped')
